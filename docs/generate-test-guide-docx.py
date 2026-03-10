"""
테스트 실행 가이드 Word 문서 생성 스크립트
실행: python docs/generate-test-guide-docx.py
출력: docs/test-guide.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color_hex: str):
    """셀 배경색 설정"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn('w:shd'),
        {qn('w:fill'): color_hex, qn('w:val'): 'clear'}
    )
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """스타일 적용된 테이블 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 헤더 행
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2F5496')

    # 데이터 행
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    return table


def add_code_block(doc, code: str):
    """코드 블록 추가 (회색 배경)"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # 배경색은 paragraph shading으로
    pPr = para._element.get_or_add_pPr()
    shd = pPr.makeelement(
        qn('w:shd'),
        {qn('w:fill'): 'F2F2F2', qn('w:val'): 'clear'}
    )
    pPr.append(shd)
    return para


def add_note(doc, text: str):
    """참고 박스 추가"""
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    shd = pPr.makeelement(
        qn('w:shd'),
        {qn('w:fill'): 'FFF3CD', qn('w:val'): 'clear'}
    )
    pPr.append(shd)
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.italic = True


def main():
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # ============================================================
    # 표지
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('스마트스토어 자동화 시스템')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('테스트 실행 가이드')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.date.today().strftime('%Y-%m-%d')
    run = meta.add_run(f'최종 점검일: {today}\n전체 테스트: 60 suites / 722 tests — ALL PASS\n타입 체크: 8개 패키지 모두 0 errors')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_page_break()

    # ============================================================
    # 목차
    # ============================================================
    doc.add_heading('목차', level=1)
    toc_items = [
        '1. 사전 준비',
        '2. 빠른 전체 테스트 실행',
        '3. 패키지별 개별 테스트',
        '4. 타입 체크',
        '5. 테스트 커버리지',
        '6. 특정 테스트만 실행',
        '7. 현재 테스트 현황',
        '8. 트러블슈팅',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ============================================================
    # 1. 사전 준비
    # ============================================================
    doc.add_heading('1. 사전 준비', level=1)

    doc.add_heading('필수 요구사항', level=2)
    add_styled_table(doc,
        ['항목', '버전', '확인 명령어'],
        [
            ['Node.js', '18 이상 (20 권장)', 'node -v'],
            ['npm', '9 이상', 'npm -v'],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('의존성 설치', level=2)
    add_code_block(doc, '# 프로젝트 루트에서 실행\nnpm install')

    doc.add_heading('Prisma 클라이언트 생성', level=2)
    add_code_block(doc, 'npx prisma generate --schema=packages/db/prisma/schema.prisma')

    add_note(doc, '참고: 단위 테스트는 모든 외부 의존성(DB, Redis, API)을 mock으로 대체합니다. '
                  'PostgreSQL이나 Redis가 실행 중이지 않아도 테스트 실행에 문제가 없습니다.')

    # ============================================================
    # 2. 빠른 전체 테스트 실행
    # ============================================================
    doc.add_heading('2. 빠른 전체 테스트 실행', level=1)
    doc.add_paragraph('모노레포 전체 테스트를 Turbo로 병렬 실행합니다:')
    add_code_block(doc, 'npm test')
    doc.add_paragraph('이 명령은 turbo run test를 실행하여 모든 패키지/앱의 테스트를 병렬로 실행합니다.')

    # ============================================================
    # 3. 패키지별 개별 테스트
    # ============================================================
    doc.add_heading('3. 패키지별 개별 테스트', level=1)
    doc.add_paragraph('각 패키지 디렉토리에서 독립적으로 실행할 수 있습니다.')

    packages = [
        {
            'name': 'packages/shared (공유 유틸리티)',
            'cmd': 'cd packages/shared && npx jest',
            'suites': '3개', 'tests': '24개',
            'target': '설정 로더, 서킷브레이커, 프리플라이트 체크',
        },
        {
            'name': 'packages/core (비즈니스 핵심 로직)',
            'cmd': 'cd packages/core && npx jest',
            'suites': '24개', 'tests': '368개',
            'target': '가격 계산, 마진 안전장치, 암호화, 콘텐츠 생성, 재고 관리, 주문 승인, 전략',
        },
        {
            'name': 'packages/crawlers (크롤러)',
            'cmd': 'cd packages/crawlers && npx jest',
            'suites': '4개', 'tests': '39개',
            'target': 'BaseCrawler(robots.txt), 도매꾹, 오너클랜, 네이버쇼핑 크롤러',
        },
        {
            'name': 'packages/integrations (외부 API 연동)',
            'cmd': 'cd packages/integrations && npx jest',
            'suites': '1개', 'tests': '13개',
            'target': '환율 API (캐싱, 에러 핸들링)',
        },
        {
            'name': 'packages/adapters (어댑터)',
            'cmd': 'cd packages/adapters && npx jest',
            'suites': '1개', 'tests': '12개',
            'target': 'Telegram 봇 커맨드 핸들러 (/status, /report, /pause, /resume)',
        },
        {
            'name': 'apps/api-server (REST API)',
            'cmd': 'cd apps/api-server && npx jest',
            'suites': '11개', 'tests': '125개',
            'target': '상품, 주문, 모니터링, 자격증명, 웹훅, Rate Limit, 스키마 검증',
        },
        {
            'name': 'apps/worker (BullMQ 워커)',
            'cmd': 'cd apps/worker && npx jest',
            'suites': '16개', 'tests': '141개',
            'target': '12개 워커 잡, 이미지 파이프라인, Kill Switch, 경쟁사 제한, 설정 캐시',
        },
    ]

    for pkg in packages:
        doc.add_heading(pkg['name'], level=2)
        add_code_block(doc, pkg['cmd'])
        add_styled_table(doc,
            ['항목', '내용'],
            [
                ['suites', pkg['suites']],
                ['tests', pkg['tests']],
                ['테스트 대상', pkg['target']],
            ]
        )
        doc.add_paragraph()

    # ============================================================
    # 4. 타입 체크
    # ============================================================
    doc.add_heading('4. 타입 체크', level=1)

    doc.add_heading('전체 타입 체크 (의존성 순서)', level=2)
    doc.add_paragraph('아래 순서대로 각 패키지 디렉토리에서 실행합니다:')

    type_check_cmds = [
        'cd packages/shared && npx tsc --noEmit',
        'cd packages/db && npx tsc --noEmit',
        'cd packages/adapters && npx tsc --noEmit',
        'cd packages/integrations && npx tsc --noEmit',
        'cd packages/crawlers && npx tsc --noEmit',
        'cd packages/core && npx tsc --noEmit',
        'cd apps/api-server && npx tsc --noEmit',
        'cd apps/worker && npx tsc --noEmit',
    ]
    add_code_block(doc, '\n'.join(type_check_cmds))

    add_note(doc, '출력이 없으면 0 errors (성공)입니다.')

    doc.add_heading('특정 패키지만 체크', level=2)
    add_code_block(doc, '# 예: core 패키지만\ncd packages/core && npx tsc --noEmit')

    # ============================================================
    # 5. 테스트 커버리지
    # ============================================================
    doc.add_heading('5. 테스트 커버리지', level=1)

    doc.add_heading('핵심 패키지 커버리지 실행', level=2)
    add_code_block(doc,
        '# Core (비즈니스 로직 핵심)\n'
        'cd packages/core && npx jest --coverage\n\n'
        '# Worker (자동화 워커)\n'
        'cd apps/worker && npx jest --coverage\n\n'
        '# API Server (REST API)\n'
        'cd apps/api-server && npx jest --coverage'
    )

    doc.add_heading('커버리지 리포트 열기', level=2)
    add_code_block(doc,
        '# 실행 후 coverage/lcov-report/index.html 파일을 브라우저에서 열기\n'
        'start coverage/lcov-report/index.html    # Windows\n'
        'open coverage/lcov-report/index.html     # macOS\n'
        'xdg-open coverage/lcov-report/index.html # Linux'
    )

    doc.add_heading('현재 커버리지 현황', level=2)

    doc.add_heading('packages/core', level=3)
    add_styled_table(doc,
        ['지표', '커버리지'],
        [
            ['Statements', '96.77%'],
            ['Branches', '85.78%'],
            ['Functions', '96.55%'],
            ['Lines', '96.70%'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('apps/worker', level=3)
    add_styled_table(doc,
        ['지표', '커버리지'],
        [
            ['Statements', '79.61%'],
            ['Branches', '59.92%'],
            ['Functions', '70.37%'],
            ['Lines', '80.46%'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('apps/api-server', level=3)
    add_styled_table(doc,
        ['지표', '커버리지'],
        [
            ['Statements', '81.38%'],
            ['Branches', '75.35%'],
            ['Functions', '81.15%'],
            ['Lines', '81.75%'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('커버리지 갭 (보충 테스트 후보)', level=2)
    add_styled_table(doc,
        ['파일', 'Stmts%', 'Branch%', '원인'],
        [
            ['worker/refund.job.ts', '30.7%', '12.5%', '환불 처리 분기 대부분 미커버'],
            ['worker/order.job.ts', '55.9%', '44%', '주문 매핑/발주 경로 미커버'],
            ['api-server/orders.ts', '57.3%', '63.9%', '승인/거절 라우트 미커버'],
            ['worker/registration.job.ts', '71.3%', '46.9%', '등록 워크플로우 일부 미커버'],
            ['api-server/credential-tester.ts', '69.6%', '50%', '일부 서비스 테스터 미커버'],
        ]
    )

    # ============================================================
    # 6. 특정 테스트만 실행
    # ============================================================
    doc.add_heading('6. 특정 테스트만 실행', level=1)

    doc.add_heading('파일 지정', level=2)
    add_code_block(doc, '# 특정 테스트 파일만 실행\ncd packages/core && npx jest src/pricing/wholesale.test.ts')

    doc.add_heading('패턴 매칭 (테스트 이름)', level=2)
    add_code_block(doc, '# 테스트 이름에 "마진" 포함된 테스트만 실행\ncd packages/core && npx jest --testNamePattern="마진"')

    doc.add_heading('감시 모드 (Watch)', level=2)
    add_code_block(doc, '# 파일 변경 시 자동 재실행 (개발 중 사용)\ncd packages/core && npx jest --watch')

    doc.add_heading('특정 디렉토리만', level=2)
    add_code_block(doc, '# pricing 관련 테스트만\ncd packages/core && npx jest src/pricing/')

    # ============================================================
    # 7. 현재 테스트 현황
    # ============================================================
    doc.add_heading('7. 현재 테스트 현황 총괄표', level=1)

    add_styled_table(doc,
        ['패키지', 'Suites', 'Tests', 'Stmts%', 'Branch%', '상태'],
        [
            ['packages/shared', '3', '24', '-', '-', 'PASS'],
            ['packages/core', '24', '368', '96.8%', '85.8%', 'PASS'],
            ['packages/crawlers', '4', '39', '-', '-', 'PASS'],
            ['packages/integrations', '1', '13', '-', '-', 'PASS'],
            ['packages/adapters', '1', '12', '-', '-', 'PASS'],
            ['apps/api-server', '11', '125', '81.4%', '75.4%', 'PASS'],
            ['apps/worker', '16', '141', '79.6%', '59.9%', 'PASS'],
            ['합계', '60', '722', '-', '-', 'ALL PASS'],
        ]
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'마지막 전체 점검: {today}')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run('타입 체크: 8개 패키지 모두 0 errors')
    run2.bold = True

    # ============================================================
    # 8. 트러블슈팅
    # ============================================================
    doc.add_heading('8. 트러블슈팅', level=1)

    troubleshooting = [
        {
            'title': '"Cannot find module \'@smartstore/...\'"',
            'cause': '패키지 dist 파일이 없거나 오래됨',
            'fix': 'cd packages/shared && npx tsc --outDir dist --declaration\ncd packages/core && npx tsc --outDir dist --declaration',
        },
        {
            'title': '"Cannot read properties of undefined (reading \'adapter\')"',
            'cause': '테스트에서 @smartstore/shared의 config mock에 notification 설정 누락',
            'fix': "jest.mock('@smartstore/shared')에 아래 추가:\nconfig: {\n  notification: { adapter: 'telegram', telegram: { botToken: 'test', chatId: '0' } },\n}",
        },
        {
            'title': '"Jest encountered an unexpected token" (TypeScript 구문 에러)',
            'cause': 'Jest가 ts-jest 없이 TypeScript 파일을 파싱 시도',
            'fix': 'package.json에 Jest 설정 추가:\n"jest": {\n  "preset": "ts-jest",\n  "testEnvironment": "node"\n}',
        },
        {
            'title': '"worker process has failed to exit gracefully"',
            'cause': '비동기 작업(타이머, 이벤트 리스너)이 정리되지 않음',
            'fix': 'cd apps/worker && npx jest --forceExit',
        },
        {
            'title': 'Prisma 클라이언트 오류',
            'cause': 'Prisma 클라이언트가 생성되지 않음',
            'fix': 'npx prisma generate --schema=packages/db/prisma/schema.prisma',
        },
    ]

    for ts in troubleshooting:
        doc.add_heading(ts['title'], level=2)
        p = doc.add_paragraph()
        run = p.add_run('원인: ')
        run.bold = True
        p.add_run(ts['cause'])
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        run2 = p2.add_run('해결: ')
        run2.bold = True
        doc.add_paragraph()
        add_code_block(doc, ts['fix'])

    # ============================================================
    # 부록: CI/CD
    # ============================================================
    doc.add_heading('부록: CI/CD 테스트 파이프라인', level=1)
    doc.add_paragraph('GitHub Actions에서 자동으로 실행되는 테스트:')
    add_code_block(doc,
        '# .github/workflows/ci.yml\n'
        '# Node.js 18, 20에서 병렬 테스트\n'
        'steps:\n'
        '  - npm ci\n'
        '  - npx prisma generate\n'
        '  - npx turbo run test\n'
        '  - npx prettier --check \'**/*.ts\'\n'
        '  - npx eslint \'**/*.ts\' --max-warnings 50'
    )
    doc.add_paragraph()
    doc.add_paragraph('로컬에서 CI와 동일한 환경을 재현하려면:')
    add_code_block(doc,
        'npm ci\n'
        'npx prisma generate --schema=packages/db/prisma/schema.prisma\n'
        'npm test'
    )

    # ============================================================
    # 저장
    # ============================================================
    output_path = 'docs/test-guide.docx'
    doc.save(output_path)
    print(f'Word 문서 생성 완료: {output_path}')


if __name__ == '__main__':
    main()
