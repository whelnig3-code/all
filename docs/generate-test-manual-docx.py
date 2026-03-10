"""
스마트스토어 자동화 시스템 — 운영 테스트 매뉴얼 Word 문서 생성
실행: python docs/generate-test-manual-docx.py
출력: docs/test-manual.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def shading(cell, color):
    pr = cell._element.get_or_add_tcPr()
    pr.append(pr.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'}))


def styled_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading(c, '2F5496')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri % 2 == 1:
                shading(c, 'D6E4F0')
    return t


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pr = p._element.get_or_add_pPr()
    pr.append(pr.makeelement(qn('w:shd'), {qn('w:fill'): 'F2F2F2', qn('w:val'): 'clear'}))


def note(doc, text):
    p = doc.add_paragraph()
    pr = p._element.get_or_add_pPr()
    pr.append(pr.makeelement(qn('w:shd'), {qn('w:fill'): 'FFF3CD', qn('w:val'): 'clear'}))
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.italic = True


def warn(doc, text):
    p = doc.add_paragraph()
    pr = p._element.get_or_add_pPr()
    pr.append(pr.makeelement(qn('w:shd'), {qn('w:fill'): 'F8D7DA', qn('w:val'): 'clear'}))
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.bold = True


def check(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(10)


def step(doc, num, text):
    p = doc.add_paragraph()
    r = p.add_run(f'Step {num}. ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    today = datetime.date.today().strftime('%Y-%m-%d')

    # ======================== 표지 ========================
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('스마트스토어 자동화 시스템')
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('운영 테스트 매뉴얼')
    r2.font.size = Pt(22)
    r2.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f'작성일: {today}\n버전: 1.0\n\n이 문서를 순서대로 따라하면\n시스템 설치부터 전체 기능 테스트까지 완료됩니다.')
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_page_break()

    # ======================== 목차 ========================
    doc.add_heading('목차', level=1)
    toc = [
        'PART 1. 환경 설치',
        '  1-1. Docker 설치 및 실행 (PostgreSQL, Redis)',
        '  1-2. 프로젝트 설치',
        '  1-3. 환경변수 설정',
        '  1-4. 데이터베이스 초기화',
        'PART 2. 계정 연동',
        '  2-1. 네이버 커머스 API 계정',
        '  2-2. 도매꾹 계정',
        '  2-3. 오너클랜 계정',
        '  2-4. 텔레그램 봇 생성',
        '  2-5. 대시보드에서 계정 등록 및 연결 테스트',
        'PART 3. 시스템 시작',
        '  3-1. 서비스 시작 순서',
        '  3-2. 정상 시작 확인',
        'PART 4. 기능 테스트',
        '  4-1. 상품 등록 테스트',
        '  4-2. 주문 처리 테스트',
        '  4-3. 배송 알림 테스트',
        '  4-4. 가격 모니터링 테스트',
        '  4-5. 재고 동기화 테스트',
        '  4-6. 주문 승인 모드 테스트',
        '  4-7. Kill Switch (자동화 제어) 테스트',
        '  4-8. 텔레그램 봇 커맨드 테스트',
        '  4-9. 대시보드 기능 테스트',
        'PART 5. 체크리스트',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(1)
        for r in p.runs:
            r.font.size = Pt(10)

    doc.add_page_break()

    # ======================== PART 1 ========================
    doc.add_heading('PART 1. 환경 설치', level=1)

    # 1-1
    doc.add_heading('1-1. Docker 설치 및 실행', level=2)
    doc.add_paragraph('이 시스템은 PostgreSQL(데이터 저장)과 Redis(작업 큐)가 필요합니다.')

    step(doc, 1, 'Docker Desktop 설치')
    doc.add_paragraph('https://www.docker.com/products/docker-desktop/ 에서 다운로드 후 설치합니다.')

    step(doc, 2, 'Docker 실행 확인')
    code(doc, 'docker --version\ndocker-compose --version')

    step(doc, 3, 'PostgreSQL + Redis 시작')
    code(doc, '# 프로젝트 루트 디렉토리에서 실행\ncd smartstore-automation\ndocker-compose up -d')

    step(doc, 4, '정상 작동 확인')
    code(doc, 'docker-compose ps\n\n# 결과 예시:\n# postgres   running   0.0.0.0:5432->5432/tcp\n# redis      running   0.0.0.0:6379->6379/tcp')

    check(doc, 'postgres 상태가 running인지 확인')
    check(doc, 'redis 상태가 running인지 확인')

    # 1-2
    doc.add_heading('1-2. 프로젝트 설치', level=2)

    step(doc, 1, 'Node.js 설치 (v18 이상, v20 권장)')
    doc.add_paragraph('https://nodejs.org/ 에서 LTS 버전 다운로드 후 설치합니다.')
    code(doc, 'node -v   # v20.x.x 이상 출력되어야 함\nnpm -v    # 9.x.x 이상 출력되어야 함')

    step(doc, 2, '의존성 설치')
    code(doc, 'cd smartstore-automation\nnpm install')
    note(doc, '참고: 처음 설치 시 2~5분 소요될 수 있습니다.')

    step(doc, 3, 'Prisma 클라이언트 생성')
    code(doc, 'npx prisma generate --schema=packages/db/prisma/schema.prisma')

    # 1-3
    doc.add_heading('1-3. 환경변수 설정', level=2)

    step(doc, 1, '.env 파일 생성')
    code(doc, 'cp .env.example .env')

    step(doc, 2, '필수 환경변수 수정')
    doc.add_paragraph('.env 파일을 열어 아래 항목을 반드시 수정합니다:')

    styled_table(doc,
        ['환경변수', '설명', '예시'],
        [
            ['DATABASE_URL', 'PostgreSQL 접속 주소', 'postgresql://user:password@localhost:5432/smartstore'],
            ['MASTER_ENCRYPTION_KEY', '암호화 키 (64자 hex)', '아래 명령어로 생성'],
            ['ADMIN_PASS', '관리자 비밀번호', '아래 명령어로 생성'],
            ['REDIS_HOST', 'Redis 주소', 'localhost'],
            ['REDIS_PORT', 'Redis 포트', '6379'],
        ]
    )

    doc.add_paragraph()
    step(doc, 3, '암호화 키 생성')
    code(doc, '# MASTER_ENCRYPTION_KEY 생성 (터미널에서 실행)\nnode -e "console.log(require(\'crypto\').randomBytes(32).toString(\'hex\'))"')
    doc.add_paragraph('출력된 64자 문자열을 .env 파일의 MASTER_ENCRYPTION_KEY에 붙여넣습니다.')

    step(doc, 4, '관리자 비밀번호 생성')
    code(doc, '# ADMIN_PASS 생성\nnode -e "console.log(require(\'crypto\').randomBytes(24).toString(\'base64\'))"')
    doc.add_paragraph('출력된 문자열을 .env 파일의 ADMIN_PASS에 붙여넣습니다.')
    warn(doc, '주의: ADMIN_PASS를 "changeme"로 두면 서버가 시작되지 않습니다.')

    # 1-4
    doc.add_heading('1-4. 데이터베이스 초기화', level=2)

    step(doc, 1, 'DB 테이블 생성')
    code(doc, 'npx prisma db push --schema=packages/db/prisma/schema.prisma')

    step(doc, 2, '테이블 생성 확인')
    code(doc, 'npx prisma studio --schema=packages/db/prisma/schema.prisma')
    doc.add_paragraph('브라우저에서 Prisma Studio가 열리면 products, orders 등 테이블이 표시됩니다.')
    check(doc, 'products 테이블이 보이는지 확인')
    check(doc, 'orders 테이블이 보이는지 확인')
    check(doc, 'system_settings 테이블이 보이는지 확인')

    doc.add_page_break()

    # ======================== PART 2 ========================
    doc.add_heading('PART 2. 계정 연동', level=1)
    doc.add_paragraph('시스템 사용을 위해 외부 서비스 계정을 준비하고 연동합니다.')

    # 2-1
    doc.add_heading('2-1. 네이버 커머스 API 계정', level=2)
    doc.add_paragraph('네이버 스마트스토어에 상품을 등록하고 주문을 관리하려면 네이버 커머스 API 계정이 필요합니다.')

    step(doc, 1, '네이버 커머스센터 접속')
    doc.add_paragraph('https://sell.smartstore.naver.com/ 에서 판매자 계정으로 로그인합니다.')

    step(doc, 2, 'API 인증 정보 발급')
    doc.add_paragraph('커머스센터 > 판매자 정보 > API 설정 페이지에서:')
    check(doc, 'Client ID 복사 → .env의 NAVER_CLIENT_ID에 입력')
    check(doc, 'Client Secret 복사 → .env의 NAVER_CLIENT_SECRET에 입력')
    check(doc, '판매자 코드(Shop ID) 복사 → .env의 NAVER_SHOP_ID에 입력')

    step(doc, 3, '필요한 값 3개')
    styled_table(doc,
        ['항목', '.env 변수명', '어디서 찾나요'],
        [
            ['Client ID', 'NAVER_CLIENT_ID', '커머스센터 > API 설정'],
            ['Client Secret', 'NAVER_CLIENT_SECRET', '커머스센터 > API 설정'],
            ['판매자 코드', 'NAVER_SHOP_ID', '커머스센터 > 판매자 정보'],
        ]
    )

    # 2-2
    doc.add_heading('2-2. 도매꾹 계정', level=2)

    step(doc, 1, '도매꾹 회원가입')
    doc.add_paragraph('https://domeggook.com/ 에서 회원가입합니다.')

    step(doc, 2, '로그인 정보 준비')
    check(doc, '아이디 → 대시보드 설정 페이지에서 입력')
    check(doc, '비밀번호 → 대시보드 설정 페이지에서 입력')

    # 2-3
    doc.add_heading('2-3. 오너클랜 계정', level=2)

    step(doc, 1, '오너클랜 회원가입')
    doc.add_paragraph('https://www.ownerclan.com/ 에서 회원가입합니다.')

    step(doc, 2, '로그인 정보 준비')
    check(doc, '아이디 → 대시보드 설정 페이지에서 입력')
    check(doc, '비밀번호 → 대시보드 설정 페이지에서 입력')

    # 2-4
    doc.add_heading('2-4. 텔레그램 봇 생성 (선택)', level=2)
    doc.add_paragraph('알림 수신과 원격 제어를 위해 텔레그램 봇을 생성합니다.')

    step(doc, 1, '텔레그램 앱에서 @BotFather 검색')
    doc.add_paragraph('텔레그램 앱을 열고 검색창에 @BotFather를 입력하여 대화를 시작합니다.')

    step(doc, 2, '새 봇 생성')
    code(doc, '/newbot\n\n# BotFather가 이름을 물으면:\nSmartStore Bot\n\n# 사용자명을 물으면:\nyour_smartstore_bot')
    doc.add_paragraph('성공하면 Bot Token이 표시됩니다. 이 토큰을 복사합니다.')

    step(doc, 3, 'Chat ID 확인')
    doc.add_paragraph('생성된 봇에게 아무 메시지를 보낸 후:')
    code(doc, '# 브라우저에서 아래 URL 접속 (TOKEN을 실제 토큰으로 교체)\nhttps://api.telegram.org/bot<TOKEN>/getUpdates')
    doc.add_paragraph('응답 JSON에서 "chat":{"id": 123456789} 값이 Chat ID입니다.')

    step(doc, 4, '환경변수 설정')
    check(doc, 'Bot Token → .env의 TELEGRAM_BOT_TOKEN에 입력')
    check(doc, 'Chat ID → .env의 TELEGRAM_CHAT_ID에 입력')

    # 2-5
    doc.add_heading('2-5. 대시보드에서 계정 등록 및 연결 테스트', level=2)
    doc.add_paragraph('모든 계정 정보를 대시보드의 설정 페이지에서 등록하고 연결을 테스트합니다.')
    note(doc, '참고: 이 단계는 PART 3에서 시스템을 시작한 후에 진행합니다.')

    step(doc, 1, '대시보드 접속: http://localhost:4000/settings')
    step(doc, 2, '판매자 유형을 선택합니다 (개인 판매자 / 사업자)')
    step(doc, 3, '필수 서비스 3개를 순서대로 설정합니다:')
    check(doc, '네이버 커머스 API → "설정" 버튼 → Client ID, Client Secret, Shop ID 입력 → 저장')
    check(doc, '도매꾹 → "설정" 버튼 → 아이디, 비밀번호 입력 → 저장')
    check(doc, '오너클랜 → "설정" 버튼 → 아이디, 비밀번호 입력 → 저장')

    step(doc, 4, '각 서비스의 "테스트" 버튼을 클릭하여 연결을 확인합니다')
    check(doc, '네이버 커머스: "OAuth 토큰 발급 성공" 메시지 확인')
    check(doc, '도매꾹: "로그인 성공" 메시지 확인')
    check(doc, '오너클랜: "로그인 성공" 메시지 확인')

    step(doc, 5, '(선택) 텔레그램 봇 설정')
    check(doc, '텔레그램 → "설정" 버튼 → Bot Token, Chat ID 입력 → 저장')
    check(doc, '"테스트" 버튼 → "봇 연결 성공: @봇이름" 확인')

    doc.add_page_break()

    # ======================== PART 3 ========================
    doc.add_heading('PART 3. 시스템 시작', level=1)

    # 3-1
    doc.add_heading('3-1. 서비스 시작 순서', level=2)
    doc.add_paragraph('3개의 서비스를 각각 별도 터미널에서 시작합니다.')
    warn(doc, '반드시 아래 순서대로 시작하세요. Docker(PostgreSQL, Redis)가 먼저 실행 중이어야 합니다.')

    step(doc, 1, '터미널 1: API 서버 시작')
    code(doc, 'cd smartstore-automation\nnpm run dev -w apps/api-server\n\n# 정상 시작 메시지:\n# Server listening on http://0.0.0.0:3100')

    step(doc, 2, '터미널 2: 워커 시작')
    code(doc, 'cd smartstore-automation\nnpm run dev -w apps/worker\n\n# 정상 시작 메시지:\n# 모든 워커 시작 완료\n# 워커 준비 완료. Phase 2~4 자동화 실행 중')

    step(doc, 3, '터미널 3: 대시보드 시작')
    code(doc, 'cd smartstore-automation\nnpm run dev -w apps/dashboard\n\n# 정상 시작 메시지:\n# ready - started server on 0.0.0.0:4000')

    # 3-2
    doc.add_heading('3-2. 정상 시작 확인', level=2)

    step(doc, 1, 'API 서버 헬스체크')
    code(doc, 'curl http://localhost:3100/\n\n# 응답 예시:\n# {"status":"ok","service":"smartstore-api","version":"1.0.0"}')
    check(doc, 'status가 "ok"인지 확인')

    step(doc, 2, '대시보드 접속')
    doc.add_paragraph('브라우저에서 http://localhost:4000 접속')
    check(doc, '대시보드 메인 화면이 표시되는지 확인')
    check(doc, '시스템 상태 섹션에서 Worker, DB, Redis가 모두 녹색(정상)인지 확인')

    step(doc, 3, '텔레그램 봇 확인 (설정한 경우)')
    doc.add_paragraph('텔레그램 앱에서 봇에게 /status 메시지를 보냅니다.')
    check(doc, '"시스템 상태" 응답이 오는지 확인')
    check(doc, 'Worker, DB, Redis가 모두 "정상"인지 확인')

    doc.add_page_break()

    # ======================== PART 4 ========================
    doc.add_heading('PART 4. 기능 테스트', level=1)
    doc.add_paragraph('아래 순서대로 각 기능을 테스트합니다. 각 테스트 항목의 "확인 사항"을 반드시 체크하세요.')

    # 4-1
    doc.add_heading('4-1. 상품 등록 테스트', level=2)
    doc.add_paragraph('도매처(도매꾹/오너클랜) 상품을 네이버 스마트스토어에 자동 등록하는 기능입니다.')

    step(doc, 1, 'API로 상품 등록 요청')
    code(doc,
        'curl -X POST http://localhost:3100/products \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{\n'
        '    "source": "domaegguk",\n'
        '    "sourceProductId": "12345",\n'
        '    "name": "테스트 USB 케이블 1m",\n'
        '    "wholesalePrice": 3000,\n'
        '    "shippingFee": 2500,\n'
        '    "naverFeeRate": 0.05,\n'
        '    "targetMarginRate": 0.30,\n'
        '    "images": ["https://example.com/image.jpg"],\n'
        '    "description": "테스트 상품입니다"\n'
        '  }\''
    )

    step(doc, 2, '응답 확인')
    check(doc, '상태코드 201 응답')
    check(doc, 'priceCalculation.salePrice에 계산된 판매가가 표시되는지 확인')
    check(doc, 'priceCalculation.marginRate가 30% 이상인지 확인')

    step(doc, 3, '등록 진행 확인')
    code(doc, '# 상품 목록 조회\ncurl http://localhost:3100/products')
    check(doc, '방금 등록한 상품이 목록에 표시되는지 확인')
    check(doc, 'status가 "pending" → "registered"로 변경되는지 확인 (워커가 처리)')

    step(doc, 4, '대시보드에서 확인')
    doc.add_paragraph('http://localhost:4000/products 접속')
    check(doc, '상품 목록에 테스트 상품이 표시되는지 확인')

    step(doc, 5, '네이버 커머스센터에서 확인')
    doc.add_paragraph('https://sell.smartstore.naver.com/ 에서 상품 관리 메뉴 확인')
    check(doc, '등록한 상품이 네이버 스마트스토어에 정상 노출되는지 확인')

    note(doc, '가격 계산 공식: 판매가 = CEIL((도매가 + 배송비) / (1 - 수수료율 - 마진율), 10원)\n'
         '예시: CEIL((3000 + 2500) / (1 - 0.05 - 0.30), 10) = CEIL(8461.5, 10) = 8,470원')

    # 4-2
    doc.add_heading('4-2. 주문 처리 테스트', level=2)
    doc.add_paragraph('네이버에서 발생한 주문을 자동으로 감지하고 처리하는 기능입니다.')

    step(doc, 1, '테스트 주문 생성')
    doc.add_paragraph('네이버 스마트스토어에서 등록된 테스트 상품을 직접 구매하거나, 테스트 주문을 생성합니다.')

    step(doc, 2, '수동 주문 폴링 트리거')
    code(doc, 'curl -X POST http://localhost:3100/orders/poll\n\n# 응답: {"message":"주문 폴링 트리거됨"}')

    step(doc, 3, '주문 목록 확인')
    code(doc, 'curl http://localhost:3100/orders')
    check(doc, '새 주문이 목록에 표시되는지 확인')
    check(doc, 'status가 "paid"인지 확인')

    step(doc, 4, '대시보드에서 확인')
    doc.add_paragraph('http://localhost:4000/orders 접속')
    check(doc, '주문 목록에 새 주문이 표시되는지 확인')

    step(doc, 5, '텔레그램 알림 확인 (설정한 경우)')
    check(doc, '"새 주문 도착" 알림 메시지가 텔레그램에 도착했는지 확인')

    # 4-3
    doc.add_heading('4-3. 배송 알림 테스트', level=2)
    doc.add_paragraph('주문에 운송장 번호를 등록하고 고객에게 배송 알림을 보내는 기능입니다.')

    step(doc, 1, '발송 처리 요청')
    code(doc,
        '# {orderId}를 실제 주문 ID로 교체\n'
        'curl -X POST http://localhost:3100/orders/{orderId}/ship \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"trackingNumber": "123456789", "courier": "CJ대한통운"}\''
    )
    check(doc, '상태코드 200 응답')
    check(doc, '주문 상태가 "preparing"으로 변경되는지 확인')

    step(doc, 2, '네이버에서 확인')
    doc.add_paragraph('커머스센터 > 주문관리에서 해당 주문의 발송 처리 상태를 확인합니다.')
    check(doc, '운송장 번호가 정상 등록되었는지 확인')

    # 4-4
    doc.add_heading('4-4. 가격 모니터링 테스트', level=2)
    doc.add_paragraph('경쟁사 가격을 자동 수집하고, 가격을 자동 조정하는 기능입니다. 매시 정각에 자동 실행됩니다.')

    step(doc, 1, '현재 등록된 상품의 가격 시뮬레이션 확인')
    code(doc, '# {productId}를 실제 상품 ID로 교체\ncurl http://localhost:3100/products/{productId}/price-simulation')
    check(doc, '마진율별 판매가 시뮬레이션 결과가 표시되는지 확인')

    step(doc, 2, '워커 로그에서 가격 모니터링 실행 확인')
    doc.add_paragraph('워커 터미널에서 매시 정각에 아래와 유사한 로그를 확인합니다:')
    code(doc, '[INFO] [price-monitor] 경쟁가 수집 완료: 5개 상품\n[INFO] [price-monitor] 가격 조정: 19,240원 → 19,230원 (경쟁가 대비 -10원)')

    step(doc, 3, '대시보드에서 가격 히스토리 확인')
    doc.add_paragraph('http://localhost:4000/products 에서 특정 상품 클릭 > 가격 이력을 확인합니다.')
    check(doc, '가격 변동 이력이 표시되는지 확인')
    check(doc, '마진율이 15% 이상 유지되는지 확인')
    warn(doc, '안전장치: 마진율이 15% 미만으로 내려가면 가격 조정이 자동 거부됩니다.')

    # 4-5
    doc.add_heading('4-5. 재고 동기화 테스트', level=2)
    doc.add_paragraph('공급처(도매꾹/오너클랜)의 재고를 10분마다 자동 동기화합니다.')

    step(doc, 1, '재고 현황 확인')
    code(doc, 'curl http://localhost:3100/inventory/status')
    check(doc, '상품별 supplierStock(공급처 재고)과 cachedStock(캐시 재고)이 표시되는지 확인')

    step(doc, 2, '수동 재고 동기화')
    code(doc, '# {productId}를 실제 상품 ID로 교체\ncurl -X POST http://localhost:3100/inventory/{productId}/sync')
    check(doc, '"재고 동기화 큐에 추가됨" 메시지 확인')

    step(doc, 3, '재동기화 후 확인')
    code(doc, 'curl http://localhost:3100/inventory/{productId}')
    check(doc, 'lastStockSync 시간이 갱신되었는지 확인')
    check(doc, '공급처 재고가 0이면 listingPaused가 true인지 확인 (자동 판매중지)')

    # 4-6
    doc.add_heading('4-6. 주문 승인 모드 테스트', level=2)
    doc.add_paragraph('주문이 들어오면 자동 처리 대신, 텔레그램으로 승인 버튼을 보내 운영자가 직접 승인/거부하는 모드입니다.')
    warn(doc, '사전 조건: 텔레그램 봇이 설정되어 있어야 합니다.')

    step(doc, 1, '승인 모드 활성화')
    code(doc,
        'curl -X POST http://localhost:3100/admin/control \\\n'
        '  -H "Authorization: Basic $(echo -n admin:비밀번호 | base64)" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"key": "ORDER_APPROVAL_MODE", "value": "true"}\''
    )
    doc.add_paragraph('또는 대시보드에서 토글 스위치를 켭니다.')

    step(doc, 2, '테스트 주문 생성 후 수동 폴링')
    code(doc, 'curl -X POST http://localhost:3100/orders/poll')

    step(doc, 3, '텔레그램 확인')
    check(doc, '텔레그램에 주문 승인 요청 메시지가 도착하는지 확인')
    check(doc, '메시지에 상품명, 판매가, 마진율, 재고 정보가 표시되는지 확인')
    check(doc, '[승인] [거부] 버튼이 표시되는지 확인')

    step(doc, 4, '승인 버튼 클릭')
    check(doc, '메시지가 "승인됨"으로 변경되는지 확인')
    check(doc, '주문 상태가 "preparing"으로 변경되는지 확인')

    step(doc, 5, '5분 타임아웃 테스트')
    doc.add_paragraph('새 주문을 생성하고 아무 버튼도 누르지 않고 5분을 기다립니다.')
    check(doc, '5분 후 자동으로 타임아웃 처리되는지 확인')
    check(doc, '텔레그램에 타임아웃 알림이 오는지 확인')

    step(doc, 6, '승인 모드 비활성화 (테스트 완료 후)')
    code(doc,
        'curl -X POST http://localhost:3100/admin/control \\\n'
        '  -H "Authorization: Basic $(echo -n admin:비밀번호 | base64)" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"key": "ORDER_APPROVAL_MODE", "value": "false"}\''
    )

    # 4-7
    doc.add_heading('4-7. Kill Switch (자동화 제어) 테스트', level=2)
    doc.add_paragraph('각 자동화 기능을 개별적으로 일시정지/재개할 수 있습니다.')

    styled_table(doc,
        ['기능', '제어 키', '기본값'],
        [
            ['가격 자동 조정', 'AUTO_PRICE_ENABLED', 'true'],
            ['주문 자동 처리', 'AUTO_ORDER_ENABLED', 'true'],
            ['배송 자동 알림', 'AUTO_SHIPPING_ENABLED', 'true'],
            ['재고 자동 동기화', 'AUTO_INVENTORY_SYNC_ENABLED', 'true'],
            ['주문 승인 모드', 'ORDER_APPROVAL_MODE', 'false'],
        ]
    )

    doc.add_paragraph()
    step(doc, 1, '가격 자동 조정 일시정지')
    code(doc,
        '# API로 일시정지\n'
        'curl -X POST http://localhost:3100/admin/control \\\n'
        '  -H "Authorization: Basic $(echo -n admin:비밀번호 | base64)" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"key": "AUTO_PRICE_ENABLED", "value": "false"}\'\n\n'
        '# 또는 텔레그램에서:\n'
        '/pause price'
    )

    step(doc, 2, '일시정지 확인')
    code(doc, '# 시스템 상태 조회\ncurl -H "Authorization: Basic ..." http://localhost:3100/admin/system')
    check(doc, '응답의 settings.AUTO_PRICE_ENABLED가 "false"인지 확인')

    step(doc, 3, '재개')
    code(doc, '# API로 재개\ncurl -X POST http://localhost:3100/admin/control \\\n'
         '  -H "Authorization: Basic ..." \\\n'
         '  -H "Content-Type: application/json" \\\n'
         '  -d \'{"key": "AUTO_PRICE_ENABLED", "value": "true"}\'\n\n'
         '# 또는 텔레그램에서:\n/resume price')
    check(doc, '설정값이 "true"로 돌아왔는지 확인')

    # 4-8
    doc.add_heading('4-8. 텔레그램 봇 커맨드 테스트', level=2)
    doc.add_paragraph('텔레그램 앱에서 봇에게 각 커맨드를 보내고 응답을 확인합니다.')

    styled_table(doc,
        ['커맨드', '기능', '예상 응답'],
        [
            ['/status', '시스템 상태 조회', 'Worker/DB/Redis 연결 상태 + 메모리'],
            ['/report', '오늘 실적 조회', '매출, 순익, 주문 수, 실패 건수'],
            ['/pending', '승인 대기 주문 목록', '대기 중인 주문 목록 또는 "없음"'],
            ['/pause price', '가격 자동화 정지', '"가격 자동화 일시정지됨"'],
            ['/resume price', '가격 자동화 재개', '"가격 자동화 재개됨"'],
            ['/pause order', '주문 자동화 정지', '"주문 자동화 일시정지됨"'],
            ['/resume order', '주문 자동화 재개', '"주문 자동화 재개됨"'],
        ]
    )

    doc.add_paragraph()
    step(doc, 1, '각 커맨드를 순서대로 텔레그램 봇에게 전송합니다')
    step(doc, 2, '각 응답이 위 표의 "예상 응답"과 일치하는지 확인합니다')
    check(doc, '/status → 시스템 상태가 정상으로 표시되는지 확인')
    check(doc, '/report → 오늘 매출/주문 수가 표시되는지 확인')
    check(doc, '/pause price → 정지 메시지 + /resume price → 재개 메시지 확인')

    # 4-9
    doc.add_heading('4-9. 대시보드 기능 테스트', level=2)
    doc.add_paragraph('http://localhost:4000 에 접속하여 각 페이지를 확인합니다.')

    step(doc, 1, '메인 대시보드 (http://localhost:4000)')
    check(doc, '오늘의 현황(매출, 주문, 순익)이 표시되는지 확인')
    check(doc, '자동화 제어 토글(가격/주문/배송)이 작동하는지 확인')
    check(doc, '시스템 상태(Worker/DB/Redis)가 모두 녹색인지 확인')

    step(doc, 2, '상품 관리 (http://localhost:4000/products)')
    check(doc, '등록한 상품 목록이 표시되는지 확인')
    check(doc, '상태별 필터(전체/대기/등록됨/실패)가 작동하는지 확인')

    step(doc, 3, '주문 관리 (http://localhost:4000/orders)')
    check(doc, '주문 목록이 표시되는지 확인')
    check(doc, '상태별 필터가 작동하는지 확인')

    step(doc, 4, '매출 리포트 (http://localhost:4000/report)')
    check(doc, '기간별(7일/30일/90일) 매출 요약이 표시되는지 확인')
    check(doc, '상위 판매 상품 목록이 표시되는지 확인')

    step(doc, 5, '서비스 설정 (http://localhost:4000/settings)')
    check(doc, '설정된 서비스에 "연결됨" 배지가 표시되는지 확인')
    check(doc, '진행 상태바가 정확한지 확인 (예: 3/5 완료)')

    doc.add_page_break()

    # ======================== PART 5 ========================
    doc.add_heading('PART 5. 최종 체크리스트', level=1)
    doc.add_paragraph('모든 테스트가 완료되었는지 확인합니다. 각 항목을 체크하세요.')

    doc.add_heading('환경 설치', level=2)
    check(doc, 'Docker(PostgreSQL, Redis) 정상 실행')
    check(doc, 'Node.js v18+ 설치 확인')
    check(doc, '의존성 설치 완료 (npm install)')
    check(doc, 'DB 테이블 생성 완료 (prisma db push)')
    check(doc, '환경변수 설정 완료 (.env)')

    doc.add_heading('계정 연동', level=2)
    check(doc, '네이버 커머스 API 연결 테스트 통과')
    check(doc, '도매꾹 로그인 테스트 통과')
    check(doc, '오너클랜 로그인 테스트 통과')
    check(doc, '(선택) 텔레그램 봇 연결 테스트 통과')

    doc.add_heading('시스템 시작', level=2)
    check(doc, 'API 서버 정상 시작 (포트 3100)')
    check(doc, '워커 정상 시작 (12개 워커 가동)')
    check(doc, '대시보드 정상 시작 (포트 4000)')
    check(doc, '시스템 상태 전체 정상 (Worker/DB/Redis)')

    doc.add_heading('기능 테스트', level=2)
    check(doc, '상품 등록: 도매처 상품 → 네이버 등록 성공')
    check(doc, '주문 처리: 주문 폴링 → DB 저장 → 알림 도착')
    check(doc, '배송 알림: 운송장 등록 → 발송 처리 성공')
    check(doc, '가격 모니터링: 경쟁가 수집 → 가격 조정 (마진 15%+ 유지)')
    check(doc, '재고 동기화: 공급처 재고 조회 → DB 갱신')
    check(doc, '주문 승인 모드: 텔레그램 승인/거부 버튼 동작')
    check(doc, 'Kill Switch: 자동화 일시정지/재개 동작')
    check(doc, '텔레그램 봇: /status, /report, /pause, /resume 응답 정상')
    check(doc, '대시보드: 모든 페이지 정상 표시')

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('모든 항목을 체크했다면 시스템이 정상 작동하는 것입니다.')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    # ======================== 저장 ========================
    out = 'docs/test-manual.docx'
    doc.save(out)
    print(f'운영 테스트 매뉴얼 생성 완료: {out}')


if __name__ == '__main__':
    main()
